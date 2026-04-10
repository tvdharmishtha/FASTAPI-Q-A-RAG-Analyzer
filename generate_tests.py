import os
from docx import Document
from fpdf import FPDF

os.makedirs('test_docs', exist_ok=True)

# 1. Text File
txt_content = """Biology: The Cell
The mitochondria is commonly known as the powerhouse of the cell. It generates most of the chemical energy needed to power the cell's biochemical reactions, storing this energy in a molecule called ATP. The human body contains trillions of cells, each harboring from 1,000 to 2,500 mitochondria."""
with open('test_docs/Biology_Cell.txt', 'w', encoding='utf-8') as f:
    f.write(txt_content)

# 2. DOCX File
doc = Document()
doc.add_heading('Geography: Continents', 0)
doc.add_paragraph("Asia is the largest continent on Earth by land area, covering approximately 44.58 million square kilometers. It is also the most populous continent, hosting roughly 60% of the world's current human population. In contrast, Antarctica is the coldest, driest, and windiest continent.")
doc.save('test_docs/Geography_Continents.docx')

# 3. PDF File
class PDF(FPDF):
    def header(self):
        self.set_font('helvetica', 'B', 15)
        self.cell(0, 10, 'Physics: Gravity', ln=True, align='C')

pdf = PDF()
pdf.add_page()
pdf.set_font('helvetica', '', 12)
pdf.multi_cell(0, 10, "Gravity is a fundamental interaction which causes mutual attraction between all things that have mass. The most famous early theory was formulated by Sir Isaac Newton in 1687, where he described gravity as a universal force. Much later, Albert Einstein proposed the General Theory of Relativity in 1915, describing gravity not as a conventional force, but as the curvature of spacetime caused by mass.")
pdf.output('test_docs/Physics_Gravity.pdf')

print("Test files generated successfully!")
